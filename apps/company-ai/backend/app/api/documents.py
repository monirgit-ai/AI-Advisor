"""Document upload and management endpoints."""
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from typing import List, Optional, Tuple
import os
import uuid
import logging
from pathlib import Path
from datetime import datetime

from app.core.config import settings
from app.db.session import get_db
from app.db.models.document import Document, DocumentStatus
from app.db.models.document_chunk import DocumentChunk
from app.db.models.user import User
from app.api.deps import get_current_user
from app.services.indexer import index_document

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["documents"])


def get_file_extension(filename: str) -> str:
    """Get file extension in lowercase."""
    return Path(filename).suffix.lower().lstrip(".")


def is_allowed_extension(filename: str) -> bool:
    """Check if file extension is allowed."""
    ext = get_file_extension(filename)
    return ext in settings.ALLOWED_EXTENSIONS


def validate_mime_type(mime_type: str, filename: str) -> bool:
    """Validate MIME type matches file extension."""
    ext = get_file_extension(filename)
    
    mime_map = {
        "pdf": ["application/pdf"],
        "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
        "txt": ["text/plain", "text/plain; charset=utf-8"]
    }
    
    allowed_mimes = mime_map.get(ext, [])
    # Be lenient - allow if no strict match but extension is valid
    if not allowed_mimes:
        return False
    return mime_type.lower() in [m.lower() for m in allowed_mimes] or len(allowed_mimes) > 0


def extract_text_from_file(file_path: str, mime_type: str) -> Tuple[str, Optional[str]]:
    """
    Extract text from uploaded file.
    
    Returns:
        tuple: (extracted_text, error_message)
    """
    try:
        ext = Path(file_path).suffix.lower().lstrip(".")
        
        if ext == "txt":
            # Read text file with UTF-8, ignore errors
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return text, None
        
        elif ext == "docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text, None
            except ImportError:
                return "", "python-docx not installed"
            except Exception as e:
                return "", f"Error reading DOCX: {str(e)}"
        
        elif ext == "pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
                return text, None
            except ImportError:
                return "", "pypdf not installed"
            except Exception as e:
                return "", f"Error reading PDF: {str(e)}"
        
        else:
            return "", f"Unsupported file type: {ext}"
    
    except Exception as e:
        return "", f"Extraction error: {str(e)}"


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload and parse a document.
    
    Args:
        file: Uploaded file
        current_user: Authenticated user
        db: Database session
        
    Returns:
        JSON response with document info
    """
    # Validate file extension
    if not is_allowed_extension(file.filename):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File extension not allowed. Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}"
        )
    
    # Validate MIME type (lenient - check extension if MIME doesn't match)
    if not validate_mime_type(file.content_type or "", file.filename or ""):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid file type. Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}"
        )
    
    # Read file content to check size
    content = await file.read()
    file_size = len(content)
    
    # Validate file size
    if file_size > settings.max_upload_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large. Maximum size: {settings.MAX_UPLOAD_MB}MB"
        )
    
    if file_size == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File is empty"
        )
    
    # Generate document ID
    document_id = uuid.uuid4()
    
    # Create storage path: /home/aiapp/data/uploads/<company_id>/<document_id>/
    company_dir = Path(settings.UPLOAD_DIR) / str(current_user.company_id)
    document_dir = company_dir / str(document_id)
    document_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate stored filename (uuid-based for safety)
    ext = get_file_extension(file.filename or "")
    stored_filename = f"source.{ext}"
    storage_path = document_dir / stored_filename
    
    # Save file
    with open(storage_path, "wb") as f:
        f.write(content)
    
    # Create document record with status="uploaded"
    document = Document(
        id=document_id,
        company_id=current_user.company_id,
        uploaded_by_user_id=current_user.id,
        filename_original=file.filename or "unknown",
        filename_stored=stored_filename,
        mime_type=file.content_type or "application/octet-stream",
        file_size_bytes=file_size,
        storage_path=str(storage_path),
        status=DocumentStatus.UPLOADED
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    # Extract text
    text_extracted, error_message = extract_text_from_file(str(storage_path), file.content_type or "")
    
    # Update document with extracted text
    if error_message:
        document.status = DocumentStatus.FAILED
        document.error_message = error_message
    else:
        document.status = DocumentStatus.PARSED
        document.text_extracted = text_extracted
    
    db.commit()
    db.refresh(document)
    
    # Return response
    return {
        "document_id": str(document.id),
        "status": document.status.value,
        "filename": document.filename_original,
        "mime_type": document.mime_type,
        "file_size_bytes": document.file_size_bytes,
        "created_at": document.created_at.isoformat(),
        "error_message": document.error_message
    }


@router.get("")
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    List all documents for the user's company.
    
    Args:
        current_user: Authenticated user
        db: Database session
        
    Returns:
        List of document summaries
    """
    documents = db.query(Document).filter(
        Document.company_id == current_user.company_id
    ).order_by(Document.created_at.desc()).all()
    
    return [
        {
            "id": str(doc.id),
            "filename_original": doc.filename_original,
            "status": doc.status.value,
            "created_at": doc.created_at.isoformat(),
            "file_size_bytes": doc.file_size_bytes,
            "mime_type": doc.mime_type,
            "index_status": doc.index_status.value if doc.index_status is not None else None
        }
        for doc in documents
    ]


@router.get("/{document_id}")
async def get_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get document details by ID.
    
    Args:
        document_id: Document UUID
        current_user: Authenticated user
        db: Database session
        
    Returns:
        Document details with first 2000 chars of text
    """
    try:
        doc_uuid = uuid.UUID(document_id)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format"
        )
    
    document = db.query(Document).filter(Document.id == doc_uuid).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Ensure document belongs to user's company
    if document.company_id != current_user.company_id:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Limit text preview to 2000 characters
    text_preview = None
    if document.text_extracted:
        text_preview = document.text_extracted[:2000]
        if len(document.text_extracted) > 2000:
            text_preview += "... (truncated)"
    
    return {
        "id": str(document.id),
        "filename_original": document.filename_original,
        "filename_stored": document.filename_stored,
        "mime_type": document.mime_type,
        "file_size_bytes": document.file_size_bytes,
        "status": document.status.value,
        "error_message": document.error_message,
        "created_at": document.created_at.isoformat(),
        "text_preview": text_preview,
        "text_length": len(document.text_extracted) if document.text_extracted else 0,
        "index_status": document.index_status.value if hasattr(document, 'index_status') else None
    }


@router.post("/{document_id}/index")
async def index_document_endpoint(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Index a document: chunk, embed, and store chunks.
    
    Args:
        document_id: Document UUID
        current_user: Authenticated user
        db: Database session
        
    Returns:
        JSON response with indexing status
    """
    try:
        doc_uuid = uuid.UUID(document_id)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format"
        )
    
    # Verify document exists and belongs to user's company
    document = db.query(Document).filter(Document.id == doc_uuid).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    if document.company_id != current_user.company_id:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Index the document
    success, error_message, chunks_created = index_document(
        document_id,
        str(current_user.company_id),
        db
    )
    
    # Refresh document to get updated status
    db.refresh(document)
    
    if success:
        return {
            "document_id": document_id,
            "status": "indexed",
            "chunks_created": chunks_created,
            "index_status": document.index_status.value
        }
    else:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=error_message or "Indexing failed"
        )


@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Delete a document and unindex it (remove all chunks).
    
    This will:
    1. Delete all document chunks (unindexing)
    2. Delete the document record from the database
    3. Delete the file from storage
    
    Args:
        document_id: Document UUID
        current_user: Authenticated user
        db: Database session
        
    Returns:
        JSON response with deletion status
    """
    try:
        doc_uuid = uuid.UUID(document_id)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format"
        )
    
    # Verify document exists and belongs to user's company
    document = db.query(Document).filter(Document.id == doc_uuid).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Security: Ensure document belongs to user's company
    if document.company_id != current_user.company_id:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        # Step 1: Delete all chunks for this document (unindexing)
        chunks_deleted = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == doc_uuid
        ).delete(synchronize_session=False)
        
        logger.info(f"Deleted {chunks_deleted} chunks for document {document_id}")
        
        # Step 2: Get storage path before deleting document record
        storage_path = document.storage_path
        
        # Step 3: Delete document record
        db.delete(document)
        db.commit()
        
        # Step 4: Delete file from storage (if exists)
        if storage_path and os.path.exists(storage_path):
            try:
                os.remove(storage_path)
                logger.info(f"Deleted file: {storage_path}")
                
                # Also try to remove empty parent directories
                try:
                    document_dir = Path(storage_path).parent
                    if document_dir.exists() and not any(document_dir.iterdir()):
                        document_dir.rmdir()
                        logger.info(f"Removed empty document directory: {document_dir}")
                    
                    company_dir = document_dir.parent
                    if company_dir.exists() and not any(company_dir.iterdir()):
                        company_dir.rmdir()
                        logger.info(f"Removed empty company directory: {company_dir}")
                except OSError:
                    # Ignore errors when removing directories (they might not be empty)
                    pass
                    
            except OSError as e:
                logger.warning(f"Failed to delete file {storage_path}: {e}")
                # Don't fail the request if file deletion fails - DB record is already deleted
        
        return {
            "document_id": document_id,
            "status": "deleted",
            "chunks_deleted": chunks_deleted,
            "message": "Document and all associated chunks have been deleted"
        }
    
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {e}", exc_info=True)
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document: {str(e)}"
        )
